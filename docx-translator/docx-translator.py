import docx
from googletrans import Translator

def get_text(file):
    doc = docx.Document(file)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def translate_text(text):
    translator = Translator()
    translated_text = translator.translate(text)
    return translated_text.text

def text_to_docx(text, filename):
    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(filename)


if __name__ == "__main__":
    file_loc = "docx-translator/notas.docx"
    text = get_text(file_loc)
    translated_text = translate_text(text)
    text_to_docx(translated_text, "test.docx")